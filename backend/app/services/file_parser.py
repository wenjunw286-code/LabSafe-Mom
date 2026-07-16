"""Async file parser for PDF, DOCX, and TXT protocol files.

Produces plain text from uploaded laboratory protocol documents.
Supports UTF-8, GBK, and Latin-1 encodings for TXT files.
"""

from __future__ import annotations

import io
from pathlib import Path

import structlog

from app.core.exceptions import FileParsingError

logger = structlog.get_logger(__name__)


class FileParser:
    """Parse laboratory protocol files (PDF, DOCX, TXT) to plain text.

    All methods are static and CPU-bound — they run in the main thread
    but are non-blocking for I/O. For very large files, consider
    running in a thread pool executor.
    """

    SUPPORTED_TYPES: set[str] = {"pdf", "docx", "txt"}

    @classmethod
    def parse(cls, file_content: bytes, filename: str) -> str:
        """Parse file content to plain text, dispatching by extension.

        Args:
            file_content: Raw file bytes.
            filename: Original filename (used to determine type).

        Returns:
            Extracted plain text.

        Raises:
            FileParsingError: If the file type is unsupported or parsing fails.
        """
        ext = Path(filename).suffix.lower().lstrip(".")
        if ext not in cls.SUPPORTED_TYPES:
            raise FileParsingError(
                f"Unsupported file type: .{ext}",
                detail={"supported": list(cls.SUPPORTED_TYPES)},
            )

        logger.info("file_parse_start", filename=filename, ext=ext, size=len(file_content))

        try:
            if ext == "txt":
                return cls._parse_txt(file_content)
            elif ext == "pdf":
                return cls._parse_pdf(file_content)
            elif ext == "docx":
                return cls._parse_docx(file_content)
        except FileParsingError:
            raise
        except Exception as exc:
            logger.error("file_parse_error", filename=filename, error=str(exc))
            raise FileParsingError(
                f"Failed to parse {ext.upper()} file: {filename}",
                detail={"error": str(exc)},
            ) from exc

        raise FileParsingError(f"Unexpected file type: {ext}")

    @staticmethod
    def _parse_txt(content: bytes) -> str:
        """Decode TXT content with encoding detection."""
        for encoding in ("utf-8", "gbk", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        # Last resort: replace invalid chars
        return content.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Extract text from PDF using pypdf."""
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        texts: list[str] = []
        for page in reader.pages:
            try:
                text = page.extract_text()
                if text:
                    texts.append(text)
            except Exception:
                # Skip pages that can't be extracted
                continue
        return "\n".join(texts)

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """Extract text from DOCX including paragraphs and tables."""
        from docx import Document

        doc = Document(io.BytesIO(content))
        texts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                texts.append(" | ".join(row_cells))

        return "\n".join(texts)
